from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs")
DOCX_PATH = OUT / "Week_9_Model_Explainer.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, size=None, color=None):
    run.bold = bold
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    style_run(run, bold=True, size=16 if level == 1 else 13, color="2E74B5" if level < 3 else "1F4D78")
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Normal"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run(text)


def add_metric_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Model", "Accuracy", "Failure Recall", "Failure Precision", "Avg Precision"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                style_run(run, bold=True)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        values = [
            row["model"],
            "" if row["accuracy"] is None else f"{row['accuracy']:.1%}",
            "" if row["failure_recall"] is None else f"{row['failure_recall']:.1%}",
            "" if row["failure_precision"] is None else f"{row['failure_precision']:.1%}",
            f"{row['average_precision']:.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            set_cell_margins(cell)
    return table


def main():
    payload = json.loads((OUT / "lab_results.json").read_text(encoding="utf-8"))
    metrics = payload["metrics"]
    baseline = next(r for r in metrics if r["model"] == "Baseline Random Forest")
    xgb = next(r for r in metrics if r["model"] == "XGBoost + Class Weighting")
    top_factors = payload["top_local_shap_factors"]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Week 9 Model Explainer: Predictive Maintenance")
    style_run(run, bold=True, size=20, color="0B2545")
    subtitle = doc.add_paragraph("Predictive Intelligence & Industrial Trust")
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("555555")
    subtitle.runs[0].font.size = Pt(11)

    add_heading(doc, "Executive Takeaway", 1)
    add_body(
        doc,
        f"This lab simulates a maintenance setting where failures are rare: only "
        f"{payload['failure_rate']:.1%} of equipment logs are failure events. The key lesson is that "
        f"accuracy alone can hide missed failures, so the workflow emphasizes recall, precision-recall curves, "
        f"and SHAP explanations operators can act on.",
    )

    add_heading(doc, "Pipeline Summary", 1)
    add_bullet(doc, f"Created {payload['dataset_shape'][0]:,} sensor records across 120 assets with an AI4I/C-MAPSS-style rare failure pattern.")
    add_bullet(doc, "Applied Min-Max Scaling to normalize vibration, temperature, torque, pressure, wear, speed, and cycle signals.")
    add_bullet(doc, "Used a leakage-aware split by asset group, so future assets do not sneak into model training.")
    add_bullet(doc, "Compared a raw Random Forest baseline, SMOTE, XGBoost class weighting, a neural sequence proxy, and K-Means high-risk segmentation.")

    add_heading(doc, "Model Results", 1)
    add_metric_table(doc, metrics)
    add_body(
        doc,
        f"The baseline Random Forest reached {baseline['accuracy']:.1%} accuracy, but its failure recall was "
        f"{baseline['failure_recall']:.1%}. XGBoost with class weighting caught "
        f"{xgb['failure_recall']:.1%} of failures and kept failure precision at {xgb['failure_precision']:.1%}. "
        "Because failures are rare, the precision-recall curve is the more honest view of operational usefulness.",
    )

    add_heading(doc, "Why the Highest-Risk Asset Was Flagged", 1)
    add_body(
        doc,
        f"For asset {payload['high_risk_asset']} at cycle {payload['high_risk_cycle']}, "
        f"the weighted XGBoost model estimated a {payload['high_risk_probability']:.1%} failure probability. "
        "SHAP explains this specific prediction by showing which signals pushed the alert upward.",
    )
    for factor in top_factors:
        direction = "toward failure" if factor["shap_contribution"] > 0 else "toward normal"
        add_bullet(
            doc,
            f"{factor['feature']}: scaled value {factor['value']:.3f}, contribution {factor['shap_contribution']:.3f} {direction}.",
        )

    add_heading(doc, "Plain-Language Explanation", 1)
    add_body(
        doc,
        "The model is like an experienced mechanic listening to an engine. It does not make the decision from one sensor alone. "
        "It combines vibration, heat, torque, pressure, and wear patterns to decide whether a machine sounds different from healthy equipment.",
    )

    add_heading(doc, "Recommended Response Path", 1)
    add_bullet(doc, "Treat high-risk predictions as triage alerts, not automatic shutdown orders.")
    add_bullet(doc, "Inspect the physical systems connected to the top SHAP drivers, such as bearings for vibration or cooling and lubrication for temperature.")
    add_bullet(doc, "Open a maintenance ticket when the same asset has repeated high-risk readings or when SHAP drivers match known failure modes.")
    add_bullet(doc, "Record technician feedback so future models learn which alerts were useful.")

    add_heading(doc, "Limitations", 1)
    add_body(
        doc,
        "This is a synthetic lab dataset designed to mirror an industrial rare-event problem. A production model would require real sensor history, "
        "verified failure labels, asset metadata, and review by maintenance experts before operational use.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Week 9 Lab")

    doc.save(DOCX_PATH)
    print(DOCX_PATH.resolve())


if __name__ == "__main__":
    main()
